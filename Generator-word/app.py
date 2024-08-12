from docx import Document
import os

def generate_certificate(template_path, output_folder, student_data):
    # Load the Word template
    doc = Document(template_path)

    for student in student_data:
        # Make a copy of the template
        doc_copy = Document(template_path)
        
        # Replace placeholders with student data
        for paragraph in doc_copy.paragraphs:
            if '{name}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{name}', student['name'])
            if '{course}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{course}', student['course'])
            if '{date}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{date}', student['date'])
        
        # Save the certificate with student's name
        output_path = os.path.join(output_folder, f"{student['name']}_certificate.docx")
        doc_copy.save(output_path)
        print(f"Generated certificate for {student['name']}")

# Example student data
student_data = [
    {'name': ' ', 'course': ' ', 'date': ' '},
    {'name': ' ', 'course': ' ', 'date': ' '},
    # Add more student data as needed
]

# Paths
template_path = r'D:\vsCode\Python\Auto\certificate-generator\Generator-word\template.docx'
output_folder = r'D:\vsCode\Python\Auto\certificate-generator\docx'

# Generate certificates
generate_certificate(template_path, output_folder, student_data)
